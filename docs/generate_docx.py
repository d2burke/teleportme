#!/usr/bin/env python3
"""Generate TeleportMe TRD as a .docx file."""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re

doc = Document()

# -- Page setup --
for section in doc.sections:
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

# -- Style configuration --
style = doc.styles['Normal']
font = style.font
font.name = 'Helvetica Neue'
font.size = Pt(10.5)
font.color.rgb = RGBColor(0x33, 0x33, 0x33)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.15

for level in range(1, 4):
    h_style = doc.styles[f'Heading {level}']
    h_style.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)
    h_style.font.name = 'Helvetica Neue'
    if level == 1:
        h_style.font.size = Pt(22)
        h_style.paragraph_format.space_before = Pt(24)
        h_style.paragraph_format.space_after = Pt(12)
    elif level == 2:
        h_style.font.size = Pt(16)
        h_style.paragraph_format.space_before = Pt(18)
        h_style.paragraph_format.space_after = Pt(8)
    else:
        h_style.font.size = Pt(13)
        h_style.paragraph_format.space_before = Pt(14)
        h_style.paragraph_format.space_after = Pt(6)


def add_table(headers, rows, col_widths=None):
    """Add a formatted table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = header
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
                run.font.size = Pt(9.5)

    # Data rows
    for row_idx, row_data in enumerate(rows):
        for col_idx, cell_text in enumerate(row_data):
            cell = table.rows[row_idx + 1].cells[col_idx]
            cell.text = str(cell_text)
            for paragraph in cell.paragraphs:
                for run in paragraph.runs:
                    run.font.size = Pt(9.5)

    doc.add_paragraph()  # spacing after table


def add_code_block(text):
    """Add a code block with monospace font."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x2D, 0x2D, 0x2D)


def add_bullet(text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        run = p.add_run(bold_prefix)
        run.bold = True
        run.font.size = Pt(10.5)
        p.add_run(text)
    else:
        p.add_run(text)


# ============================================================
# TITLE PAGE
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('TeleportMe')
run.font.size = Pt(36)
run.bold = True
run.font.color.rgb = RGBColor(0x0A, 0x0A, 0x0A)

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('Technical Requirements Document')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_paragraph()

meta = doc.add_paragraph()
meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
meta.add_run('Version 1.0  |  February 12, 2026  |  MVP (In Development)').font.size = Pt(11)

doc.add_paragraph()
doc.add_paragraph()

meta2 = doc.add_paragraph()
meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = meta2.add_run('Audience: ')
run.font.size = Pt(11)
run = meta2.add_run('Principal Architects, Engineering Teams, Agent Teams')
run.font.size = Pt(11)
run.italic = True

doc.add_page_break()

# ============================================================
# 1. EXECUTIVE SUMMARY
# ============================================================
doc.add_heading('1. Executive Summary', level=1)

doc.add_paragraph(
    'TeleportMe is an iOS-native city relocation discovery app. Users describe their lifestyle '
    'preferences, and the system recommends cities worldwide that match their priorities across '
    'cost, climate, culture, and job market dimensions. The app combines algorithmic scoring with '
    'GPT-4o-mini refinement to produce personalized city matches with AI-generated insights.'
)

doc.add_paragraph(
    'Architecture class: Client-heavy SwiftUI app with Supabase BaaS (Backend-as-a-Service) and a '
    'single Edge Function for AI-powered report generation.'
)

doc.add_paragraph(
    'Current state: MVP with 55 seeded cities, 17 score categories per city, onboarding flow, '
    'report generation, saved cities, and offline-first caching.'
)

# ============================================================
# 2. SYSTEM ARCHITECTURE
# ============================================================
doc.add_heading('2. System Architecture Overview', level=1)

doc.add_paragraph(
    'The system is divided into two primary tiers: the iOS client and the Supabase cloud backend. '
    'The client handles all presentation, state management, and local caching. The backend provides '
    'authentication (GoTrue), auto-generated REST API (PostgREST), and a Deno-based Edge Function '
    'that orchestrates algorithmic scoring and GPT refinement for report generation.'
)

add_code_block(
    'iOS Client (SwiftUI)\n'
    '  Presentation    ->  12 feature views + 11 reusable components\n'
    '  Coordination    ->  AppCoordinator (@Observable)\n'
    '  Service Layer   ->  AuthService, CityService, ReportService, SavedCitiesService\n'
    '  Persistence     ->  CacheManager (JSON files, TTL, per-user scoping)\n'
    '       |\n'
    '  HTTPS (REST API)\n'
    '       |\n'
    'Supabase Cloud\n'
    '  Auth (GoTrue)   ->  JWT sessions, email/password\n'
    '  PostgREST       ->  Auto-generated REST from PostgreSQL schema\n'
    '  Edge Functions   ->  generate-report (Deno + OpenAI GPT-4o-mini)\n'
    '  PostgreSQL 17   ->  8 tables, RLS policies, pg_trgm fuzzy search'
)

# ============================================================
# 3. CLIENT ARCHITECTURE
# ============================================================
doc.add_heading('3. Client Architecture', level=1)

doc.add_heading('3.1 Platform & Dependencies', level=2)

add_table(
    ['Attribute', 'Value'],
    [
        ['Platform', 'iOS 17+ (SwiftUI)'],
        ['Language', 'Swift 6 (strict concurrency)'],
        ['State management', '@Observable (Observation framework)'],
        ['Navigation', 'NavigationStack with typed NavigationPath'],
        ['Package manager', 'Swift Package Manager'],
        ['External dependencies', 'supabase-swift v2.33.1 (Auth, PostgREST, Functions, Realtime, Storage)'],
        ['Color scheme', 'Dark only (enforced)'],
        ['Bundle ID', 'app.TeleportMe'],
    ]
)

doc.add_heading('3.2 Layer Architecture', level=2)

doc.add_paragraph().add_run('Presentation Layer').bold = True
doc.add_paragraph(
    'All 12 feature views are pure SwiftUI. No UIKit wrapping except for UINavigationBarAppearance '
    'configuration in the app entry point. Views access shared state through '
    '@Environment(AppCoordinator.self) and create local bindings with @Bindable when two-way binding is needed.'
)

doc.add_paragraph().add_run('Coordination Layer (AppCoordinator)').bold = True
doc.add_paragraph(
    'Single @Observable coordinator manages screen routing (splash/onboarding/main), '
    'onboarding navigation via NavigationPath with OnboardingStep enum (8 steps), all shared '
    'onboarding state, service lifecycle, cache restoration on relaunch, and sign-out with full state cleanup.'
)

doc.add_paragraph().add_run('Service Layer').bold = True
doc.add_paragraph(
    'Four domain services (Auth, City, Report, SavedCities) plus SupabaseManager singleton. '
    'All are @Observable and owned by the coordinator. Every network call is wrapped with '
    'withTimeout() for deadline-based cancellation using structured concurrency (withThrowingTaskGroup).'
)

doc.add_paragraph().add_run('Persistence Layer (CacheManager)').bold = True
doc.add_paragraph(
    'JSON-file-based caching in Application Support/TeleportMe/Cache/ with global vs per-user '
    'directories, TTL-based staleness checking, atomic writes, and automatic corrupt file deletion.'
)

doc.add_heading('3.3 File Inventory', level=2)

add_table(
    ['Directory', 'File', 'Purpose'],
    [
        ['App Root', 'TeleportMeApp.swift', 'Entry point, nav bar config, DEBUG_SCREEN support'],
        ['Core/Navigation', 'AppCoordinator.swift', 'Central state + navigation coordinator'],
        ['Core/Navigation', 'RootView.swift', 'Root router, OnboardingFlowView, MainTabView'],
        ['Core/Models', 'Models.swift', 'All Codable data models (14 types)'],
        ['Core/Services', 'SupabaseManager.swift', 'Supabase client singleton'],
        ['Core/Services', 'AuthService.swift', 'Auth: sign up, sign in, sign out, profile CRUD'],
        ['Core/Services', 'CityService.swift', 'City fetching, search, scores (cached)'],
        ['Core/Services', 'ReportService.swift', 'Report generation + past reports (cached)'],
        ['Core/Services', 'SavedCitiesService.swift', 'Save/unsave with optimistic updates (cached)'],
        ['Core/Utilities', 'CacheManager.swift', 'Disk cache with TTL + per-user scoping'],
        ['Core/Utilities', 'AsyncHelpers.swift', 'withTimeout(), Collection[safe:]'],
        ['Core/Design', 'TeleportTheme.swift', 'Colors, typography, spacing, corner radii'],
        ['Core/Design', 'Components.swift', '8 reusable UI components'],
        ['Core/Design', 'PreviewHelpers.swift', 'Mock data + PreviewContainer'],
        ['Core/Debug', 'DevModeView.swift', 'DEBUG-only dev tools panel'],
        ['Features/Onboarding', 'SplashView.swift', 'Landing screen'],
        ['Features/Onboarding', 'NameInputView.swift', 'Collect user name'],
        ['Features/Onboarding', 'SignUpView.swift', 'Email/password registration + validation'],
        ['Features/Onboarding', 'SignInView.swift', 'Email/password login'],
        ['Features/Onboarding', 'StartTypeView.swift', 'Choose start type'],
        ['Features/Discovery', 'CitySearchView.swift', 'Search + trending cities'],
        ['Features/Discovery', 'CityBaselineView.swift', 'Selected city metrics display'],
        ['Features/Discovery', 'PreferencesView.swift', '4 preference sliders'],
        ['Features/Results', 'GeneratingView.swift', 'Loading animation during report gen'],
        ['Features/Results', 'RecommendationsView.swift', 'Match cards + comparison metrics'],
        ['Features/Results', 'ReportDetailView.swift', 'Individual report detail'],
        ['Features/Tabs', 'DiscoverView.swift', 'Browse/search tab'],
        ['Features/Tabs', 'SavedView.swift', 'Saved cities tab'],
        ['Features/Tabs', 'MapView.swift', 'Map visualization (placeholder)'],
        ['Features/Tabs', 'ProfileView.swift', 'User profile + settings + sign out'],
    ]
)

doc.add_heading('3.4 Navigation Model', level=2)

doc.add_paragraph(
    'The app has three screen states: splash, onboarding, and main. The onboarding flow uses a '
    'NavigationStack with an 8-step enum (NameInput through Recommendations). The generating-to-recommendations '
    'transition is a replace (removeLast + append) to prevent back-swiping to the loading screen. '
    'The main screen uses a TabView with 4 tabs: Discover, Saved, Map, Profile.'
)

doc.add_heading('3.5 Caching Strategy', level=2)

add_table(
    ['Cache Key', 'Scope', 'TTL', 'Write Trigger'],
    [
        ['cities', 'Global', '24 hours', 'fetchAllCities() network success'],
        ['cityScores', 'Global', '24 hours', 'getCityWithScores() network success'],
        ['savedCities(userId)', 'Per-user', '5 minutes', 'Every save/unsave (optimistic + confirmed)'],
        ['currentReport(userId)', 'Per-user', 'None (persistent)', 'generateReport() success'],
        ['pastReports(userId)', 'Per-user', '1 hour', 'loadReports() network success'],
        ['preferences(userId)', 'Per-user', 'None (persistent)', 'savePreferences()'],
        ['selectedCity(userId)', 'Per-user', 'None (persistent)', 'selectCity()'],
    ]
)

doc.add_paragraph().add_run('Cache-then-network pattern:').bold = True
add_bullet('Load from disk cache and populate @Observable properties instantly')
add_bullet('If cache is fresh (within TTL), skip network call')
add_bullet('Otherwise fetch from Supabase in background')
add_bullet('On success: update properties + write-through to cache')
add_bullet('On failure: keep cached data, only show error if cache was empty')

doc.add_paragraph().add_run('Optimistic update pattern (SavedCitiesService):').bold = True
add_bullet('Update local state + write cache immediately')
add_bullet('Send network request in background')
add_bullet('On success: replace optimistic entry with server response + write cache')
add_bullet('On failure: rollback local state + write cache (rollback state)')

doc.add_heading('3.6 Error Handling', level=2)

doc.add_paragraph(
    'Every service follows a consistent pattern: withTimeout() wraps all Supabase calls '
    '(10-30s depending on operation). isLoading/loadError observable properties drive three-state '
    'UI: loading spinner, error with retry button, or content. Errors only surface when no cached '
    'data is available.'
)

doc.add_heading('3.7 Design System', level=2)

doc.add_paragraph().add_run('Color Palette (dark theme only):').bold = True

add_table(
    ['Token', 'Hex', 'Usage'],
    [
        ['background', '#0A0A0A', 'App background (near black)'],
        ['surface', '#1A1A1A', 'Card backgrounds'],
        ['surfaceElevated', '#222222', 'Active/selected states'],
        ['border', '#2A2A2A', 'Dividers, inactive progress bars'],
        ['accent', '#D4FF00', 'Primary CTA, highlights (lime/chartreuse)'],
        ['textPrimary', '#FFFFFF', 'Headings, primary text'],
        ['textSecondary', '#8E8E93', 'Labels, descriptions'],
        ['textTertiary', '#636366', 'Placeholders, metadata'],
    ]
)

doc.add_paragraph(
    'Typography: System font with 7 semantic styles (heroTitle 40px through caption 12px). '
    'Score values use .rounded design variant. '
    'Spacing: 4px increments (xs:4, sm:8, md:16, lg:24, xl:32, xxl:48). '
    'Corner radii: small:8, medium:12, card:16, large:20, pill:28. '
    'Reusable components: TeleportButton (3 styles), CardView, ScoreBar, ComparisonBar, '
    'MetricCard, TrendingChip, SectionHeader, CityHeroImage, PreferenceSliderCard.'
)

# ============================================================
# 4. SERVICE-SIDE ARCHITECTURE
# ============================================================
doc.add_heading('4. Service-Side Architecture', level=1)

doc.add_heading('4.1 Supabase Configuration', level=2)

add_table(
    ['Service', 'Details'],
    [
        ['Project ref', 'REDACTED_PROJECT_REF'],
        ['PostgreSQL', 'v17'],
        ['Auth', 'GoTrue (email/password, JWT, 1h token expiry)'],
        ['API', 'PostgREST (auto-generated REST)'],
        ['Edge Functions', 'Deno v1 runtime, oneshot policy'],
        ['Realtime', 'Enabled (not yet used)'],
        ['Storage', 'Enabled (not yet used)'],
    ]
)

doc.add_heading('4.2 Database Schema', level=2)

doc.add_paragraph().add_run('Reference data (public read):').bold = True

doc.add_paragraph(
    'cities: 55 rows with id (text PK slug), name, full_name, country, continent, lat/long, '
    'population, teleport_city_score, summary, image_url, timestamps.'
)
doc.add_paragraph(
    'city_scores: 935 rows (17 categories per city). Unique constraint on (city_id, category). '
    'Categories: Housing, Cost of Living, Startups, Venture Capital, Travel Connectivity, Commute, '
    'Business Freedom, Safety, Healthcare, Education, Environmental Quality, Economy, Taxation, '
    'Internet Access, Leisure & Culture, Tolerance, Outdoors. All scores on 0-10 scale.'
)
doc.add_paragraph(
    'vibe_tags + city_vibe_tags: Tag taxonomy with junction table (schema exists, not yet wired to client).'
)

doc.add_paragraph().add_run('User data (private per user via RLS):').bold = True

doc.add_paragraph(
    'profiles: Extends auth.users via FK. Stores name, email, current_city_id, avatar_url. '
    'Auto-created by trigger on user signup.'
)
doc.add_paragraph(
    'user_preferences: Start type, 4 preference sliders (0-10), vibe tag selection, dealbreakers.'
)
doc.add_paragraph(
    'city_reports: Stores generated reports with preferences snapshot (JSONB) and results array (JSONB). '
    'References current city.'
)
doc.add_paragraph(
    'saved_cities: User-city junction with unique constraint. Used for heart/bookmark feature.'
)
doc.add_paragraph(
    'engagement_events: Event tracking table (schema exists, not yet used by client).'
)

doc.add_heading('4.3 Row Level Security', level=2)

add_table(
    ['Table', 'SELECT', 'INSERT', 'UPDATE', 'DELETE'],
    [
        ['profiles', 'Own only', 'Own (via trigger)', 'Own only', '-'],
        ['cities', 'Public', '-', '-', '-'],
        ['city_scores', 'Public', '-', '-', '-'],
        ['vibe_tags / city_vibe_tags', 'Public', '-', '-', '-'],
        ['user_preferences', 'Own only', 'Own only', 'Own only', 'Own only'],
        ['city_reports', 'Own only', 'Own only', '-', '-'],
        ['saved_cities', 'Own only', 'Own only', '-', 'Own only'],
        ['engagement_events', 'Own only', 'Own only', '-', '-'],
    ]
)

doc.add_heading('4.4 Database Functions & Indexes', level=2)

doc.add_paragraph().add_run('Functions:').bold = True
add_bullet('search_cities(search_term): ', 'Fuzzy search via pg_trgm with 0.15 similarity threshold. Returns top 20.')
add_bullet('handle_new_user(): ', 'Trigger that auto-creates profile row on auth.users insert.')
add_bullet('update_updated_at(): ', 'Generic trigger for updated_at timestamp management.')

doc.add_paragraph().add_run('Key indexes:').bold = True
add_bullet('GIN trigram indexes on cities.name and cities.full_name (fuzzy search)')
add_bullet('B-tree indexes on all foreign keys and frequently queried columns')
add_bullet('Descending index on city_reports.created_at for recent-first ordering')

doc.add_heading('4.5 Edge Function: generate-report', level=2)

doc.add_paragraph().add_run('Runtime: ').bold = True
doc.paragraphs[-1].add_run('Deno v1 (Supabase Edge Functions)')

doc.add_paragraph().add_run('Dependencies: ').bold = True
doc.paragraphs[-1].add_run('deno.land/std HTTP server, @supabase/supabase-js@2.39.3, openai@v4.38.0')

doc.add_paragraph().add_run('Algorithm:').bold = True
add_bullet('Authenticate user from JWT in Authorization header')
add_bullet('Fetch current city + its 17 scores from database')
add_bullet('Fetch all candidate cities (excluding current) + all scores')
add_bullet('Algorithmic scoring: weighted sum of 4 primary categories + 5 bonus categories. Clamp to 40-99%.')
add_bullet('Take top 8 candidates for GPT refinement')
add_bullet('GPT-4o-mini: picks 3-5 best matches, writes personalized 1-sentence insights per city')
add_bullet('Validate GPT response: parse JSON, filter to valid candidate IDs only')
add_bullet('Fallback: if GPT parse fails, use top 4 algorithmic picks with generic insights')
add_bullet('Build comparison metrics for 5 display categories (cost, climate, culture, economy, commute)')
add_bullet('Save report to city_reports table (if authenticated user)')
add_bullet('Return structured response with report_id, current_city info, and matches array')

doc.add_paragraph().add_run('Timeouts: ').bold = True
doc.paragraphs[-1].add_run('30s client-side via withTimeout(). Edge function uses Supabase default timeout.')

# ============================================================
# 5. DATA FLOW
# ============================================================
doc.add_heading('5. Key Data Flows', level=1)

doc.add_heading('5.1 App Relaunch (Returning User)', level=2)
add_bullet('RootView.onAppear triggers AppCoordinator.checkExistingSession()')
add_bullet('AuthService.checkSession() validates Supabase JWT (10s timeout)')
add_bullet('On auth success: restoreCachedState(userId) loads preferences, currentReport, selectedCity from disk')
add_bullet('CityService.fetchAllCities() loads from disk cache (instant) then refreshes from network if stale')
add_bullet('goToMain() transitions to MainTabView with cached data already rendered')

doc.add_heading('5.2 Report Generation', level=2)
add_bullet('User adjusts preference sliders and taps "Find My Cities"')
add_bullet('Coordinator saves preferences (cache + Supabase upsert) and pushes GeneratingView')
add_bullet('GeneratingView triggers generateReport() which invokes the edge function (30s timeout)')
add_bullet('Edge function scores cities algorithmically, refines with GPT, saves to DB, returns matches')
add_bullet('Client caches response and replaces GeneratingView with RecommendationsView')

doc.add_heading('5.3 Save/Unsave City', level=2)
add_bullet('User taps heart icon, triggering SavedCitiesService.toggleSave()')
add_bullet('Local state updated optimistically (instant UI response)')
add_bullet('Cache written immediately with optimistic state')
add_bullet('Network request sent in background')
add_bullet('On failure: local state rolled back and cache updated with rollback state')

# ============================================================
# 6. SECURITY
# ============================================================
doc.add_heading('6. Security Analysis', level=1)

doc.add_heading('6.1 Current State', level=2)

add_table(
    ['Area', 'Status', 'Risk Level'],
    [
        ['API key storage', 'Hardcoded in SupabaseManager.swift', 'HIGH'],
        ['Authentication', 'Email/password via Supabase GoTrue', 'Medium'],
        ['Authorization', 'RLS on all tables', 'Good'],
        ['Network transport', 'HTTPS to Supabase cloud', 'Good'],
        ['Local storage', 'JSON files in Application Support (no encryption)', 'Medium'],
        ['Edge function auth', 'JWT + service role key server-side', 'Good'],
        ['OpenAI API key', 'Server-side env var only', 'Good'],
        ['Password policy', 'Min 6 chars, no complexity requirements', 'Medium'],
    ]
)

doc.add_heading('6.2 Recommendations', level=2)
add_bullet('Move API keys to Config.xcconfig excluded from version control, or use build-time injection')
add_bullet('Enable email confirmation in Supabase auth config')
add_bullet('Add cache encryption for sensitive per-user data (CryptoKit AES-GCM)')
add_bullet('Implement certificate pinning for production builds')
add_bullet('Add rate limiting to edge function')
add_bullet('Rotate the anon key (already committed to source control)')

# ============================================================
# 7. RISKS & TECH DEBT
# ============================================================
doc.add_heading('7. Risks & Technical Debt', level=1)

doc.add_heading('7.1 High Priority', level=2)

add_table(
    ['Risk', 'Impact', 'Mitigation'],
    [
        ['Hardcoded Supabase credentials', 'Key exposure in source control', 'Move to .xcconfig or Keychain; rotate current key'],
        ['No test coverage', 'Regressions undetectable', 'Add unit tests for CacheManager, services, coordinator'],
        ['Single edge function', 'Feature loss if function fails', 'Add health check; implement client-side algo fallback'],
        ['GPT dependency', 'API outage blocks report gen', 'Algorithmic fallback exists but generic insights'],
        ['No CI/CD pipeline', 'Manual builds only', 'Set up Xcode Cloud or GitHub Actions'],
    ]
)

doc.add_heading('7.2 Medium Priority', level=2)

add_table(
    ['Risk', 'Impact', 'Mitigation'],
    [
        ['No image caching', 'City images re-downloaded each session', 'URLCache config or SDWebImage/Kingfisher'],
        ['No pagination', '55 cities fully in memory', 'Add cursor-based pagination before 500+ cities'],
        ['Cache corruption', 'Empty state until network fetch', 'Add migration versioning to cache format'],
        ['No analytics pipeline', 'engagement_events table unused', 'Implement event tracking for key actions'],
        ['Map tab is placeholder', 'Incomplete feature visible to users', 'Implement MapKit or remove tab from MVP'],
        ['Vibe tags unused', 'Schema exists but no client integration', 'Wire up or remove from schema'],
    ]
)

doc.add_heading('7.3 Future Considerations', level=2)

add_bullet('Push notifications (config exists, not wired)')
add_bullet('Deep linking / Universal Links')
add_bullet('Accessibility audit (Dynamic Type, VoiceOver)')
add_bullet('Localization (English only, strings not externalized)')
add_bullet('A/B testing framework')
add_bullet('Realtime subscriptions for live sync')
add_bullet('User-uploaded content via Supabase Storage')

# ============================================================
# 8. PERFORMANCE
# ============================================================
doc.add_heading('8. Performance Characteristics', level=1)

add_table(
    ['Operation', 'Latency', 'Notes'],
    [
        ['App launch to main (cached)', '< 200ms', 'Auth session + cache restore'],
        ['App launch to main (cold)', '2-5s', 'Auth + network fetch cities'],
        ['City search (local)', '< 10ms', 'In-memory filter of 55 cities'],
        ['City search (server)', '200-500ms', 'pg_trgm fuzzy search'],
        ['Report generation', '5-15s', 'Edge function + GPT-4o-mini'],
        ['Save/unsave city', '< 50ms', 'Optimistic update; network in background'],
        ['Cache file read', '< 5ms', 'Small JSON files (< 200KB total)'],
    ]
)

doc.add_paragraph(
    'Memory: Total app memory estimated < 50MB (excluding images). '
    'Disk: Cache directory < 500KB total. Per-user cache < 100KB.'
)

# ============================================================
# 9. ROADMAP
# ============================================================
doc.add_heading('9. Improvements Roadmap', level=1)

doc.add_heading('Phase 1: Stabilization (Current)', level=2)
add_bullet('Run seed data against remote database')
add_bullet('Unpause Supabase project and verify all endpoints')
add_bullet('End-to-end test: sign up, generate report, relaunch, verify cache')
add_bullet('Fix or remove MapView placeholder')
add_bullet('Add basic unit tests for CacheManager and service layer')

doc.add_heading('Phase 2: Production Readiness', level=2)
add_bullet('Move API credentials to secure config')
add_bullet('Enable email confirmation')
add_bullet('Add CI/CD pipeline (Xcode Cloud or GitHub Actions)')
add_bullet('Implement engagement event tracking')
add_bullet('Add image caching layer')
add_bullet('Error reporting integration (Sentry or similar)')
add_bullet('App Store submission preparation')

doc.add_heading('Phase 3: Feature Expansion', level=2)
add_bullet('Vibe tags integration (browse cities by vibe)')
add_bullet('Interactive map with city pins and filters')
add_bullet('Social features (share reports, compare with friends)')
add_bullet('Push notifications for new city data')
add_bullet('Edit profile (name, avatar)')
add_bullet('Report comparison over time')
add_bullet('City detail deep-dive view')

doc.add_heading('Phase 4: Scale', level=2)
add_bullet('Expand to 200+ cities')
add_bullet('Automated city data scoring pipeline')
add_bullet('Cursor-based pagination')
add_bullet('CDN for city images')
add_bullet('Consider SwiftData for richer local storage')
add_bullet('WebSocket real-time sync for saved cities')
add_bullet('Multi-language support')

# ============================================================
# 10. DEV INFRASTRUCTURE
# ============================================================
doc.add_heading('10. Development Infrastructure', level=1)

doc.add_heading('10.1 Build Configuration', level=2)

add_code_block(
    'DEVELOPER_DIR=/Applications/Xcode.app/Contents/Developer \\\n'
    '  xcodebuild -scheme TeleportMe \\\n'
    '  -destination \'platform=iOS Simulator,id=7AE50D1F-...\' build'
)

doc.add_paragraph('Simulator: iPhone 17 Pro, iOS 26.2')

doc.add_heading('10.2 Debug Screens', level=2)

doc.add_paragraph(
    'Set SIMCTL_CHILD_DEBUG_SCREEN environment variable to jump directly to specific screens:'
)

add_table(
    ['Value', 'Target Screen'],
    [
        ['cityBaseline', 'CityBaselineView with mock selected city'],
        ['preferences', 'PreferencesView with mock state'],
        ['generating', 'GeneratingView (loading animation)'],
        ['recommendations', 'RecommendationsView with mock report'],
        ['main', 'MainTabView with mock report'],
    ]
)

doc.add_heading('10.3 Preview System', level=2)

doc.add_paragraph(
    'All views include #Preview blocks using PreviewContainer, which injects AppCoordinator with '
    'mock data, applies dark color scheme, and provides sample cities, scores, matches, and reports '
    'via PreviewHelpers.'
)

doc.add_heading('10.4 Local Supabase Development', level=2)

add_code_block(
    'supabase start           # Start local instance\n'
    'supabase db reset        # Run migrations + seed.sql\n'
    'supabase functions deploy generate-report\n'
    '\n'
    '# Required env vars:\n'
    '# OPENAI_API_KEY, SUPABASE_URL, SUPABASE_SERVICE_ROLE_KEY, SUPABASE_ANON_KEY'
)

# ============================================================
# 11. GLOSSARY
# ============================================================
doc.add_heading('11. Glossary', level=1)

add_table(
    ['Term', 'Definition'],
    [
        ['City Score', 'A 0-10 rating for a city across one of 17 quality-of-life categories'],
        ['Match Percent', 'Algorithmic + AI-refined score (40-99%) indicating city-preference alignment'],
        ['Edge Function', 'Serverless Deno function hosted on Supabase infrastructure'],
        ['RLS', 'Row Level Security - PostgreSQL policies restricting data access per-user'],
        ['TTL', 'Time To Live - Duration before cached data is considered stale'],
        ['Write-Through', 'Cache written simultaneously with remote persistence'],
        ['Optimistic Update', 'Local state updated before network confirmation, with rollback on failure'],
        ['pg_trgm', 'PostgreSQL extension for fuzzy text matching using trigram similarity'],
        ['GoTrue', "Supabase's authentication service (JWT-based)"],
        ['PostgREST', "Supabase's auto-generated REST API from PostgreSQL schema"],
    ]
)

# ============================================================
# FOOTER
# ============================================================
doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('This document is maintained alongside the codebase and should be updated as architecture evolves.')
run.italic = True
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)

# Save
output_path = '/Users/burke/Documents/TeleportMe/docs/TeleportMe_TRD.docx'
doc.save(output_path)
print(f'Saved to {output_path}')
